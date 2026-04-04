import csv
import json
import re
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from fastapi import UploadFile
from langchain_chroma import Chroma
from langchain_core.documents import Document

from ..config import (
    BACKEND_ROOT,
    SUPABASE_STORAGE_BUCKET,
    USE_SUPABASE_VECTORS,
    get_chroma_collection_name,
    get_primary_db_path,
)
from ..models import DocumentChunkList, DocumentPipeline, PartitionCounts
from .rag_service import build_embeddings_for_provider, get_rag_service
from .supabase_service import get_supabase_service, new_uuid, utc_now


STAGE_DEFINITIONS = [
    ("upload", "Upload to S3"),
    ("queued", "Queued"),
    ("partitioning", "Partitioning"),
    ("chunking", "Chunking"),
    ("summarisation", "Summarisation"),
    ("vectorization", "Vectorization & Storage"),
]


def _slugify_filename(filename: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", filename).strip("-")
    return cleaned or "document"


def _extract_paragraphs(text: str, page: int | None = None) -> list[dict[str, Any]]:
    blocks = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    if not blocks and text.strip():
        blocks = [text.strip()]

    elements: list[dict[str, Any]] = []
    for block in blocks:
        kind = "header" if len(block) <= 80 and "\n" not in block else "text"
        elements.append({"kind": kind, "content": block, "page": page})
    return elements


def _summarize_text(text: str) -> str:
    normalized = " ".join(text.split())
    if len(normalized) <= 180:
        return normalized

    sentences = re.split(r"(?<=[.!?])\s+", normalized)
    summary = " ".join(sentences[:2]).strip()
    if summary and len(summary) >= 80:
        return summary[:220].rstrip()

    return normalized[:220].rstrip()


class IngestionService:
    def __init__(self) -> None:
        self.temp_root = BACKEND_ROOT / "data" / "tmp"
        self.temp_root.mkdir(parents=True, exist_ok=True)
        self.supabase = get_supabase_service()

    def create_document(
        self, upload: UploadFile, provider: str, user_id: str
    ) -> DocumentPipeline:
        document_id = new_uuid()
        safe_name = _slugify_filename(upload.filename or "document")
        file_bytes = upload.file.read()
        content_type = upload.content_type or "application/octet-stream"
        created_at = utc_now()
        storage_key = f"{user_id}/{document_id}/{safe_name}"

        self.supabase.upload_file(storage_key, file_bytes, content_type)
        document = self.supabase.insert_document(
            {
                "id": document_id,
                "user_id": user_id,
                "filename": upload.filename or safe_name,
                "storage_key": storage_key,
                "storage_bucket": SUPABASE_STORAGE_BUCKET,
                "file_size": len(file_bytes),
                "content_type": content_type,
                "provider": provider,
                "status": "queued",
                "current_stage": "queued",
                "error_message": None,
                "created_at": created_at,
                "updated_at": created_at,
            }
        )
        self.supabase.insert_document_event(
            document_id=document_id,
            stage_key="upload",
            stage_label="Upload to S3",
            status="completed",
            detail=f"Uploaded {upload.filename or safe_name} to Supabase Storage.",
            progress_current=1,
            progress_total=1,
        )
        self.supabase.insert_document_event(
            document_id=document_id,
            stage_key="queued",
            stage_label="Queued",
            status="running",
            detail="Document queued for processing.",
            progress_current=0,
            progress_total=0,
        )
        return self._build_pipeline_from_row(document)

    def list_documents_for_user(self, user_id: str | None) -> list[DocumentPipeline]:
        if not user_id:
            return []
        documents = self.supabase.list_documents(user_id)
        events_by_document = self._events_by_document([document["id"] for document in documents])
        chunk_counts = self._chunks_by_document([document["id"] for document in documents])
        return [
            self.supabase.build_document_pipeline(
                document,
                events_by_document.get(document["id"], []),
                chunk_counts.get(document["id"], []),
            )
            for document in documents
        ]

    def get_document(self, document_id: str, user_id: str | None = None) -> DocumentPipeline:
        if not user_id:
            raise KeyError(document_id)
        document = self.supabase.get_document(document_id, user_id)
        if not document:
            raise KeyError(document_id)
        return self._build_pipeline_from_row(document)

    def get_chunks(
        self, document_id: str, user_id: str | None = None
    ) -> DocumentChunkList:
        if not user_id:
            raise KeyError(document_id)
        document = self.supabase.get_document(document_id, user_id)
        if not document:
            raise KeyError(document_id)
        chunks = self.supabase.list_document_chunks(document_id)
        return self.supabase.build_document_chunk_list(document, chunks)

    def delete_document(self, document_id: str, user_id: str | None = None) -> None:
        if not user_id:
            raise KeyError(document_id)
        document = self.supabase.get_document(document_id, user_id)
        if not document:
            raise KeyError(document_id)

        chunks = self.supabase.list_document_chunks(document_id)
        chunk_ids = [chunk.get("vector_id") for chunk in chunks if chunk.get("vector_id")]
        if chunk_ids and not USE_SUPABASE_VECTORS:
            embeddings, _, _ = build_embeddings_for_provider(document["provider"])
            vector_store = Chroma(
                persist_directory=str(get_primary_db_path()),
                embedding_function=embeddings,
                collection_name=get_chroma_collection_name(document["provider"]),
                collection_metadata={"hnsw:space": "cosine"},
            )
            vector_store.delete(ids=chunk_ids)

        self.supabase.delete_file(document["storage_key"])
        self.supabase.delete_document(document_id, user_id)
        get_rag_service.cache_clear()

    def process_document(self, document_id: str) -> None:
        document = None
        try:
            document = self._find_document_any_owner(document_id)
            if not document:
                return

            self._mark_stage(document, "queued", "completed", "Waiting slot cleared")
            self._mark_stage(
                document,
                "partitioning",
                "running",
                "Processing and extracting text, images, and tables",
            )
            elements, counts = self._partition_document(document)
            self._mark_stage(
                document,
                "partitioning",
                "completed",
                f"Partitioned into {len(elements)} atomic elements",
                progress_current=len(elements),
                progress_total=len(elements),
            )

            self._mark_stage(document, "chunking", "running", "Creating semantic chunks")
            chunks = self._chunk_elements(elements)
            self._mark_stage(
                document,
                "chunking",
                "completed",
                f"Chunked {len(elements)} elements into {len(chunks)} chunks",
                progress_current=len(chunks),
                progress_total=len(chunks),
            )

            self._mark_stage(document, "summarisation", "running", "Creating chunk summaries")
            self._summarize_chunks(document, chunks)
            self._mark_stage(
                document,
                "summarisation",
                "completed",
                f"Summarised {len(chunks)} chunks",
                progress_current=len(chunks),
                progress_total=len(chunks),
            )

            self._mark_stage(
                document,
                "vectorization",
                "running",
                "Embedding and storing chunks",
                progress_current=0,
                progress_total=len(chunks),
            )
            self._vectorize_document(document, chunks)
            self._mark_stage(
                document,
                "vectorization",
                "completed",
                f"Stored {len(chunks)} chunks in collection '{get_chroma_collection_name(document['provider'])}'",
                progress_current=len(chunks),
                progress_total=len(chunks),
            )

            self.supabase.update_document(
                document["id"],
                document["user_id"],
                {
                    "status": "ready",
                    "current_stage": "view_chunks",
                    "error_message": None,
                },
            )
        except Exception as exc:
            if document:
                self.supabase.update_document(
                    document["id"],
                    document["user_id"],
                    {
                        "status": "failed",
                        "current_stage": document.get("current_stage", "queued"),
                        "error_message": str(exc),
                    },
                )
                self.supabase.insert_document_event(
                    document_id=document["id"],
                    stage_key=document.get("current_stage", "queued"),
                    stage_label=self._stage_label(document.get("current_stage", "queued")),
                    status="failed",
                    detail=str(exc),
                )

    def _partition_document(
        self, document: dict[str, Any]
    ) -> tuple[list[dict[str, Any]], dict[str, int]]:
        suffix = Path(document["filename"]).suffix.lower()
        file_bytes = self.supabase.download_file(document["storage_key"])
        with NamedTemporaryFile(
            suffix=suffix or ".bin",
            dir=self.temp_root,
            delete=False,
        ) as handle:
            handle.write(file_bytes)
            temp_path = Path(handle.name)

        try:
            if suffix == ".pdf":
                return self._partition_pdf(temp_path)
            if suffix == ".csv":
                return self._partition_csv(temp_path)
            if suffix == ".docx":
                return self._partition_docx(temp_path)
            return self._partition_text(temp_path)
        finally:
            temp_path.unlink(missing_ok=True)

    def _partition_pdf(self, path: Path) -> tuple[list[dict[str, Any]], dict[str, int]]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        elements: list[dict[str, Any]] = []
        counts = PartitionCounts()
        for page_index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            page_elements = _extract_paragraphs(text, page=page_index)
            elements.extend(page_elements)
            counts.text_sections += sum(1 for item in page_elements if item["kind"] == "text")
            counts.titles_headers += sum(1 for item in page_elements if item["kind"] == "header")

        return elements, counts.model_dump()

    def _partition_text(self, path: Path) -> tuple[list[dict[str, Any]], dict[str, int]]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        elements = _extract_paragraphs(text)
        counts = PartitionCounts(
            text_sections=sum(1 for item in elements if item["kind"] == "text"),
            titles_headers=sum(1 for item in elements if item["kind"] == "header"),
        )
        return elements, counts.model_dump()

    def _partition_csv(self, path: Path) -> tuple[list[dict[str, Any]], dict[str, int]]:
        rows: list[str] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                rows.append(" | ".join(cell.strip() for cell in row))

        content = "\n".join(rows)
        elements = [{"kind": "table", "content": content, "page": None}] if content else []
        counts = PartitionCounts(tables=1 if elements else 0)
        return elements, counts.model_dump()

    def _partition_docx(self, path: Path) -> tuple[list[dict[str, Any]], dict[str, int]]:
        from docx import Document as DocxDocument

        document = DocxDocument(str(path))
        text = "\n\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        elements = _extract_paragraphs(text)
        counts = PartitionCounts(
            text_sections=sum(1 for item in elements if item["kind"] == "text"),
            titles_headers=sum(1 for item in elements if item["kind"] == "header"),
        )
        return elements, counts.model_dump()

    def _chunk_elements(self, elements: list[dict[str, Any]], target_size: int = 1400) -> list[dict[str, Any]]:
        chunks: list[dict[str, Any]] = []
        current_parts: list[dict[str, Any]] = []
        current_length = 0

        for element in elements:
            content = element["content"].strip()
            if not content:
                continue

            projected = current_length + len(content) + 2
            if current_parts and projected > target_size:
                chunks.append(self._build_chunk(len(chunks), current_parts))
                current_parts = []
                current_length = 0

            current_parts.append(element)
            current_length += len(content) + 2

        if current_parts:
            chunks.append(self._build_chunk(len(chunks), current_parts))

        return chunks

    def _build_chunk(self, chunk_index: int, elements: list[dict[str, Any]]) -> dict[str, Any]:
        content = "\n\n".join(item["content"] for item in elements)
        first_page = next((item["page"] for item in elements if item["page"]), None)
        kind = "table" if any(item["kind"] == "table" for item in elements) else "text"
        return {
            "id": f"chunk-{chunk_index + 1}",
            "chunk_index": chunk_index,
            "kind": kind,
            "page": first_page,
            "char_count": len(content),
            "content": content,
            "summary": None,
        }

    def _summarize_chunks(self, document: dict[str, Any], chunks: list[dict[str, Any]]) -> None:
        total = len(chunks)
        for index, chunk in enumerate(chunks, start=1):
            chunk["summary"] = _summarize_text(chunk["content"])
            self.supabase.insert_document_event(
                document_id=document["id"],
                stage_key="summarisation",
                stage_label="Summarisation",
                status="running",
                detail="Processing chunks and creating concise summaries",
                progress_current=index,
                progress_total=total,
            )

    def _vectorize_document(self, document: dict[str, Any], chunks: list[dict[str, Any]]) -> None:
        provider = document["provider"]
        embeddings, _, _ = build_embeddings_for_provider(provider)

        db_chunks: list[dict[str, Any]] = []
        now = utc_now()
        texts = [chunk["content"] for chunk in chunks]
        embedding_vectors = embeddings.embed_documents(texts) if texts else []

        vector_documents: list[Document] = []
        vector_ids: list[str] = []
        if not USE_SUPABASE_VECTORS:
            db_path = get_primary_db_path()
            collection_name = get_chroma_collection_name(provider)
            vector_store = Chroma(
                persist_directory=str(db_path),
                embedding_function=embeddings,
                collection_name=collection_name,
                collection_metadata={"hnsw:space": "cosine"},
            )
        else:
            collection_name = "supabase_pgvector"

        for index, chunk in enumerate(chunks):
            vector_id = f"{document['id']}:{chunk['chunk_index']}"
            raw_payload = json.dumps(
                {
                    "raw_text": chunk["content"],
                    "summary": chunk["summary"],
                }
            )
            if not USE_SUPABASE_VECTORS:
                vector_documents.append(
                    Document(
                        page_content=chunk["content"],
                        metadata={
                            "source": document["filename"],
                            "page": chunk["page"],
                            "chunk_index": chunk["chunk_index"],
                            "kind": chunk["kind"],
                            "original_content": raw_payload,
                        },
                    )
                )
                vector_ids.append(vector_id)
            db_chunks.append(
                {
                    "id": new_uuid(),
                    "document_id": document["id"],
                    "chunk_index": chunk["chunk_index"],
                    "kind": chunk["kind"],
                    "page_number": chunk["page"],
                    "char_count": chunk["char_count"],
                    "content": chunk["content"],
                    "summary": chunk["summary"],
                    "vector_id": vector_id,
                    "embedding": embedding_vectors[index] if index < len(embedding_vectors) else None,
                    "created_at": now,
                }
            )

        if vector_documents and not USE_SUPABASE_VECTORS:
            vector_store.add_documents(documents=vector_documents, ids=vector_ids)
        self.supabase.replace_document_chunks(document["id"], db_chunks)
        get_rag_service.cache_clear()

    def _mark_stage(
        self,
        document: dict[str, Any],
        key: str,
        status: str,
        detail: str | None = None,
        progress_current: int = 0,
        progress_total: int = 0,
    ) -> None:
        next_status = document["status"]
        if status == "failed":
            next_status = "failed"
        elif key == "queued" and status in {"running", "completed"}:
            next_status = "queued"
        elif status == "completed" and key == "vectorization":
            next_status = "processing"
        elif status == "running":
            next_status = "processing"

        self.supabase.update_document(
            document["id"],
            document["user_id"],
            {
                "current_stage": key,
                "status": next_status,
            },
        )
        document["current_stage"] = key
        document["status"] = next_status
        self.supabase.insert_document_event(
            document_id=document["id"],
            stage_key=key,
            stage_label=self._stage_label(key),
            status=status,
            detail=detail,
            progress_current=progress_current,
            progress_total=progress_total,
        )

    def _events_by_document(self, document_ids: list[str]) -> dict[str, list[dict[str, Any]]]:
        events = self.supabase.list_document_events(document_ids)
        grouped: dict[str, list[dict[str, Any]]] = {}
        for event in events:
            grouped.setdefault(event["document_id"], []).append(event)
        return grouped

    def _chunks_by_document(self, document_ids: list[str]) -> dict[str, list[dict[str, Any]]]:
        grouped: dict[str, list[dict[str, Any]]] = {}
        for document_id in document_ids:
            grouped[document_id] = self.supabase.list_document_chunks(document_id)
        return grouped

    def _build_pipeline_from_row(self, document: dict[str, Any]) -> DocumentPipeline:
        events_by_document = self._events_by_document([document["id"]])
        chunks_by_document = self._chunks_by_document([document["id"]])
        return self.supabase.build_document_pipeline(
            document,
            events_by_document.get(document["id"], []),
            chunks_by_document.get(document["id"], []),
        )

    def _find_document_any_owner(self, document_id: str) -> dict[str, Any] | None:
        rows = self.supabase.list_documents_for_processing(document_id)
        return rows[0] if rows else None

    def _stage_label(self, key: str) -> str:
        for stage_key, label in STAGE_DEFINITIONS:
            if stage_key == key:
                return label
        return key


_SERVICE = IngestionService()


def get_ingestion_service() -> IngestionService:
    return _SERVICE
